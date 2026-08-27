from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "reports" / "tradingview-community-analysis-2026-08-27.md"
OUT = ROOT / "reports" / "TradingView_Community_Analysis_2026-08-27.docx"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.first_child_found_in('w:tcW')
    if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')

def set_table_geometry(table):
    width_map = {
        3: [2200, 1800, 5360],
        4: [900, 3300, 1800, 3360],
        5: [700, 2850, 1450, 1450, 2910],
        6: [700, 2350, 1450, 1350, 1250, 2260],
    }
    widths = width_map.get(len(table.columns), [9360 // len(table.columns)] * len(table.columns))
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr; tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa')
    ind = tblPr.first_child_found_in('w:tblInd')
    if ind is None: ind = OxmlElement('w:tblInd'); tblPr.append(ind)
    ind.set(qn('w:w'), '120'); ind.set(qn('w:type'), 'dxa')
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i] if i < len(widths) else widths[-1])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs: p.paragraph_format.space_after = Pt(2)

def add_hyperlink(paragraph, url):
    run = paragraph.add_run(url); run.font.color.rgb = RGBColor(0x05,0x63,0xC1); run.underline = True

def main():
    doc = Document(); sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
    for name,size,color,before,after in [('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6),('Heading 3',12,'1F4D78',8,4)]:
        s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    footer = sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run('TradingView and Community Analysis').font.size=Pt(9)
    lines = SRC.read_text(encoding='utf-8').splitlines(); i=0
    title = doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_after=Pt(4); rr=title.add_run('TradingView and Community Analysis'); rr.bold=True; rr.font.size=Pt(23); rr.font.color.rgb=RGBColor(0x0B,0x25,0x45)
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_after=Pt(18); sr=sub.add_run('Data through 27 August 2026'); sr.font.size=Pt(12); sr.font.color.rgb=RGBColor(0x55,0x55,0x55)
    while i < len(lines):
        line=lines[i]
        if not line.strip(): i+=1; continue
        if line.startswith('# '): i+=1; continue
        if line.startswith('## '): doc.add_heading(line[3:], level=1); i+=1; continue
        if line.startswith('### '): doc.add_heading(line[4:], level=2); i+=1; continue
        if line.startswith('|') and i+1<len(lines) and lines[i+1].startswith('|---'):
            headers=[x.strip() for x in line.strip('|').split('|')]; i+=2; data=[]
            while i<len(lines) and lines[i].startswith('|'):
                data.append([x.strip() for x in lines[i].strip('|').split('|')]); i+=1
            table=doc.add_table(rows=1, cols=len(headers)); table.style='Table Grid'; set_table_geometry(table)
            for j,h in enumerate(headers): table.cell(0,j).text=h; set_cell_shading(table.cell(0,j),'E8EEF5')
            for row in data:
                cells=table.add_row().cells
                for j,val in enumerate(row): cells[j].text=val
            set_table_geometry(table); doc.add_paragraph().paragraph_format.space_after=Pt(2); continue
        if line.startswith('- '):
            p=doc.add_paragraph(style='List Bullet'); p.add_run(line[2:]); i+=1; continue
        if line.startswith('Positive/negative counts') or line.startswith('Author counts'):
            p=doc.add_paragraph(); r=p.add_run(line); r.italic=True; r.font.color.rgb=RGBColor(0x55,0x55,0x55); i+=1; continue
        p=doc.add_paragraph();
        # Keep the report readable while preserving markdown emphasis.
        for part in re.split(r'(\*\*.*?\*\*)', line):
            if part.startswith('**') and part.endswith('**'):
                r=p.add_run(part[2:-2]); r.bold=True
            else: p.add_run(part)
        i+=1
    OUT.parent.mkdir(exist_ok=True); doc.save(OUT); print(OUT)

if __name__=='__main__': main()
